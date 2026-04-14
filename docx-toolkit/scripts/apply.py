#!/usr/bin/env python3
"""Apply JSON tree back to docx. Reads original docx as template for styles,
then rebuilds document body from JSON."""
import argparse
import json
import os
import sys
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def make_hyperlink(paragraph, text, url):
    """Add a hyperlink run to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    run_el = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    rstyle = OxmlElement('w:rStyle')
    rstyle.set(qn('w:val'), 'Hyperlink')
    rpr.append(rstyle)
    run_el.append(rpr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    run_el.append(t)
    hyperlink.append(run_el)
    paragraph._element.append(hyperlink)


def add_runs(paragraph, runs):
    """Add runs to a paragraph with formatting."""
    for r in runs:
        text = r.get("text", "")
        if "hyperlink" in r:
            make_hyperlink(paragraph, text, r["hyperlink"])
        else:
            # Split on \n and \t to restore w:br and w:tab
            import re
            parts = re.split(r'(\n|\t)', text)
            run = paragraph.add_run()
            if r.get("bold"):
                run.bold = True
            if r.get("italic"):
                run.italic = True
            for part in parts:
                if part == '\n':
                    run.add_break()
                elif part == '\t':
                    run._element.append(OxmlElement('w:tab'))
                elif part:
                    run.add_text(part)


_list_numid_cache = {}

def _get_list_numid(doc, style):
    """Get or create a numId for bullet/number lists."""
    if style in _list_numid_cache:
        return _list_numid_cache[style]
    fmt = 'bullet' if style == 'bullet' else 'decimal'
    numbering_part = doc.part.numbering_part
    num_xml = numbering_part._element
    # Find an existing abstractNum with matching format
    for absnum in num_xml.findall(qn('w:abstractNum')):
        lvl0 = absnum.find(qn('w:lvl'))
        if lvl0 is not None:
            nf = lvl0.find(qn('w:numFmt'))
            if nf is not None and nf.get(qn('w:val')) == fmt:
                aid = absnum.get(qn('w:abstractNumId'))
                # Find num referencing this abstractNum
                for num in num_xml.findall(qn('w:num')):
                    ref = num.find(qn('w:abstractNumId'))
                    if ref is not None and ref.get(qn('w:val')) == aid:
                        nid = int(num.get(qn('w:numId')))
                        _list_numid_cache[style] = nid
                        return nid
    return 1  # fallback


def _add_image_from_template(doc, template_doc, filename):
    """Copy an image from template docx into the new doc."""
    if template_doc is None:
        p = doc.add_paragraph()
        p.add_run("[Image: {}]".format(filename))
        return
    # Reuse existing relationship in doc if image already present
    for rel_id, rel in doc.part.rels.items():
        if 'image' in rel.reltype and os.path.basename(rel.target_ref) == filename:
            p = doc.add_paragraph()
            _insert_inline_image(p, rel_id, filename)
            return
    # Copy from template
    for rel_id, rel in template_doc.part.rels.items():
        if 'image' in rel.reltype and os.path.basename(rel.target_ref) == filename:
            new_rid = doc.part.relate_to(rel.target_part, rel.reltype)
            p = doc.add_paragraph()
            _insert_inline_image(p, new_rid, filename)
            return
    p = doc.add_paragraph()
    p.add_run("[Image: {}]".format(filename))


def _insert_inline_image(paragraph, r_id, filename):
    """Insert minimal inline image XML into paragraph."""
    from docx.shared import Emu
    cx, cy = Emu(5486400), Emu(3200400)  # ~5.76x3.36 inches default
    inline_xml = (
        '<wp:inline xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
        ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
        ' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        ' xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        '<wp:extent cx="{cx}" cy="{cy}"/>'
        '<wp:docPr id="1" name="{name}"/>'
        '<a:graphic><a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        '<pic:pic><pic:nvPicPr><pic:cNvPr id="0" name="{name}"/><pic:cNvPicPr/></pic:nvPicPr>'
        '<pic:blipFill><a:blip r:embed="{rid}"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill>'
        '<pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>'
        '<a:prstGeom prst="rect"><a:avLst/></a:prstGeom></pic:spPr>'
        '</pic:pic></a:graphicData></a:graphic></wp:inline>'
    ).format(cx=cx, cy=cy, name=filename, rid=r_id)
    from lxml import etree
    inline_el = etree.fromstring(inline_xml)
    drawing = OxmlElement('w:drawing')
    drawing.append(inline_el)
    run_el = OxmlElement('w:r')
    run_el.append(drawing)
    paragraph._element.append(run_el)


def add_content_blocks(doc, content, template_doc=None):
    """Add content blocks (paragraphs, tables, lists, images) to document."""
    for block in content:
        btype = block.get("type")
        if btype == "paragraph":
            p = doc.add_paragraph()
            style_name = block.get("style")
            if style_name:
                try:
                    p.style = doc.styles[style_name]
                except KeyError:
                    pass
            add_runs(p, block.get("runs", []))
        elif btype == "table":
            rows = block.get("rows", [])
            if not rows:
                continue
            ncols = max(len(row) for row in rows)
            tbl = doc.add_table(rows=0, cols=ncols)
            for row_data in rows:
                row = tbl.add_row()
                for i, cell_data in enumerate(row_data):
                    if i < ncols:
                        cell = row.cells[i]
                        # Clear default paragraph
                        cell.paragraphs[0].clear()
                        add_runs(cell.paragraphs[0], cell_data.get("runs", []))
        elif btype == "list":
            for item in block.get("items", []):
                p = doc.add_paragraph()
                style_name = item.get("style")
                if style_name:
                    try:
                        p.style = doc.styles[style_name]
                    except KeyError:
                        pass
                add_runs(p, item.get("runs", []))
                # Inject numPr directly — style-based lists don't always generate it
                ppr = p._element.find(qn('w:pPr'))
                if ppr is None:
                    ppr = OxmlElement('w:pPr')
                    p._element.insert(0, ppr)
                numpr = OxmlElement('w:numPr')
                ilvl = OxmlElement('w:ilvl')
                ilvl.set(qn('w:val'), '0')
                numid = OxmlElement('w:numId')
                numid.set(qn('w:val'), str(_get_list_numid(doc, block.get("style", "bullet"))))
                numpr.append(ilvl)
                numpr.append(numid)
                ppr.append(numpr)
        elif btype == "image":
            filename = block.get("filename", "")
            _add_image_from_template(doc, template_doc, filename)


def build_doc(data, template_path):
    """Build a new docx from JSON tree using template for styles."""
    template_doc = Document(template_path)
    doc = Document(template_path)

    # Clear existing body content
    body = doc.element.body
    for child in list(body):
        tag = child.tag
        if tag in (qn('w:p'), qn('w:tbl')):
            body.remove(child)

    nodes = data["nodes"]

    def render_node(nid):
        node = nodes[nid]
        heading = node.get("heading", "")
        level = node.get("level", 0)
        content = node.get("content", [])
        children = node.get("children", [])

        if level > 0 and heading:
            doc.add_heading(heading, level=level)

        add_content_blocks(doc, content, template_doc)

        for cid in children:
            render_node(cid)

    for rid in data["root"]:
        render_node(rid)

    return doc


def main():
    parser = argparse.ArgumentParser(description="Apply JSON tree to docx")
    parser.add_argument("template", help="Template/original .docx file (used for styles)")
    parser.add_argument("json_file", help="JSON tree file")
    parser.add_argument("-o", "--output", required=True, help="Output .docx file")
    args = parser.parse_args()

    data = load_json(args.json_file)
    doc = build_doc(data, args.template)
    doc.save(args.output)
    print("Saved to {}".format(args.output))


def load_json(path):
    with open(path, 'r', encoding='utf-8') as f:
        return json.load(f)


if __name__ == "__main__":
    main()

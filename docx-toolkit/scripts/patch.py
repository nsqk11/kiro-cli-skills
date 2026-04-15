#!/usr/bin/env python3
"""Surgical patch: apply change instructions to docx by modifying XML in-place.

Usage::

    python3 patch.py input.docx instructions.json [-o output.docx]

If ``-o`` is omitted, ``input.docx`` is overwritten.

Instructions format (JSON array)::

    [
        {"op": "update_text",    "idx": 4, "run": 0, "text": "New text"},
        {"op": "update_runs",    "idx": 4, "runs": [{"text": "Hello", "bold": true}]},
        {"op": "update_cell",    "idx": 5, "row": 0, "col": 1, "runs": [...]},
        {"op": "rename_heading", "idx": 3, "text": "New Heading"},
        {"op": "add_row",        "idx": 5, "after_row": 2, "cells": [...]},
        {"op": "delete_row",     "idx": 5, "row": 3},
        {"op": "delete",         "idx": 7},
        {"op": "add_after",      "idx": 4, "runs": [...], "clone_style_from": 4},
        {"op": "add_table_after","idx": 4, "rows": [[...]], "clone_style_from": 5},
        {"op": "insert_image",   "idx": 4, "image_path": "fig.png", "width_cm": 15},
        {"op": "move",           "idx": 10, "after": 4},
        {"op": "reply_comment",  "comment_id": 1, "text": "Done.", "author": "AI"}
    ]
"""

import argparse
import json
import os
from copy import deepcopy
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

from lxml import etree
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.shape import CT_Inline, CT_Picture
from docx.shared import Cm

# WordprocessingDrawing namespace for shape ID lookup.
_WP_NS = ('http://schemas.openxmlformats.org/drawingml/2006/'
           'wordprocessingDrawing')

# Type aliases.
Body = etree._Element
Instr = Dict[str, Any]


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _clear_runs(p_el: etree._Element) -> None:
    """Remove all ``w:r``, ``w:hyperlink``, and ``w:sdt`` children from *p_el*.

    Keeps ``w:pPr`` intact.
    """
    removable = {qn('w:r'), qn('w:hyperlink'), qn('w:sdt')}
    for child in list(p_el):
        if child.tag in removable:
            p_el.remove(child)


def _build_run_el(run_dict: Dict[str, Any]) -> etree._Element:
    """Build a ``w:r`` element from a run dict.

    Args:
        run_dict: Dict with ``text`` and optional ``bold``, ``italic``,
            ``hidden`` keys.

    Returns:
        A new ``w:r`` lxml element.
    """
    run_el = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    has_props = False
    if run_dict.get('bold'):
        rpr.append(OxmlElement('w:b'))
        has_props = True
    if run_dict.get('italic'):
        rpr.append(OxmlElement('w:i'))
        has_props = True
    if run_dict.get('hidden'):
        rpr.append(OxmlElement('w:vanish'))
        has_props = True
    if has_props:
        run_el.append(rpr)

    text = run_dict.get('text', '')
    parts: List[Tuple[str, Optional[str]]] = []
    buf: List[str] = []
    for ch in text:
        if ch == '\n':
            if buf:
                parts.append(('t', ''.join(buf)))
                buf = []
            parts.append(('br', None))
        elif ch == '\t':
            if buf:
                parts.append(('t', ''.join(buf)))
                buf = []
            parts.append(('tab', None))
        else:
            buf.append(ch)
    if buf:
        parts.append(('t', ''.join(buf)))

    for kind, val in parts:
        if kind == 't':
            t_el = OxmlElement('w:t')
            t_el.set(qn('xml:space'), 'preserve')
            t_el.text = val
            run_el.append(t_el)
        elif kind == 'br':
            run_el.append(OxmlElement('w:br'))
        elif kind == 'tab':
            run_el.append(OxmlElement('w:tab'))
    return run_el


def _build_hyperlink_el(
    part: Any,
    run_dict: Dict[str, Any],
) -> etree._Element:
    """Build a ``w:hyperlink`` element."""
    url = run_dict['hyperlink']
    rel_type = (
        'http://schemas.openxmlformats.org/officeDocument/'
        '2006/relationships/hyperlink'
    )
    r_id = part.relate_to(url, rel_type, is_external=True)
    hl_el = OxmlElement('w:hyperlink')
    hl_el.set(qn('r:id'), r_id)
    run_el = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    rstyle = OxmlElement('w:rStyle')
    rstyle.set(qn('w:val'), 'Hyperlink')
    rpr.append(rstyle)
    run_el.append(rpr)
    t_el = OxmlElement('w:t')
    t_el.set(qn('xml:space'), 'preserve')
    t_el.text = run_dict.get('text', '')
    run_el.append(t_el)
    hl_el.append(run_el)
    return hl_el


def _inject_runs(
    p_el: etree._Element,
    runs: List[Dict[str, Any]],
    part: Any,
) -> None:
    """Append run elements to *p_el* from a runs list."""
    for rd in runs:
        if rd.get('hyperlink'):
            p_el.append(_build_hyperlink_el(part, rd))
        else:
            p_el.append(_build_run_el(rd))


def _get_cell_el(
    tbl_el: etree._Element,
    row: int,
    col: int,
) -> etree._Element:
    """Return the ``w:tc`` element at *row*, *col*.

    Raises:
        ValueError: If row or col is out of range.
    """
    trs = tbl_el.findall(qn('w:tr'))
    if row >= len(trs):
        raise ValueError(
            f'Row {row} out of range (table has {len(trs)} rows)')
    tcs = trs[row].findall(qn('w:tc'))
    if col >= len(tcs):
        raise ValueError(
            f'Col {col} out of range (row {row} has {len(tcs)} cols)')
    return tcs[col]


def _next_shape_id(body: Body) -> int:
    """Find next available shape ID in the document body."""
    used: Set[int] = set()
    for doc_pr in body.findall(f'.//{{{_WP_NS}}}docPr'):
        id_val = doc_pr.get('id')
        if id_val:
            used.add(int(id_val))
    return max(used, default=0) + 1


def _get_comments_part(doc: Document) -> Any:
    """Return the comments XML part, or ``None``."""
    for rel in doc.part.rels.values():
        if 'comments' in rel.reltype:
            return rel.target_part
    return None


def _next_comment_id(comments_el: etree._Element) -> int:
    """Find next available comment ID."""
    used: Set[int] = set()
    for comment in comments_el.findall(qn('w:comment')):
        cid = comment.get(qn('w:id'))
        if cid:
            used.add(int(cid))
    return max(used, default=0) + 1


# ---------------------------------------------------------------------------
# Modification ops (don't change body length)
# ---------------------------------------------------------------------------


def _op_update_text(
    body: Body, idx: int, instr: Instr, part: Any,
) -> None:
    """Change text of a specific run, preserving all formatting."""
    el = body[idx]
    run_idx = instr['run']
    all_runs: List[etree._Element] = []
    for child in el:
        if child.tag == qn('w:r'):
            all_runs.append(child)
        elif child.tag == qn('w:hyperlink'):
            all_runs.extend(child.findall(qn('w:r')))
    if run_idx >= len(all_runs):
        raise ValueError(
            f'Run index {run_idx} out of range '
            f'(element has {len(all_runs)} runs)')
    target = all_runs[run_idx]
    for tag in ('w:t', 'w:br', 'w:tab'):
        for child in target.findall(qn(tag)):
            target.remove(child)
    t_el = OxmlElement('w:t')
    t_el.set(qn('xml:space'), 'preserve')
    t_el.text = instr['text']
    target.append(t_el)


def _op_update_runs(
    body: Body, idx: int, instr: Instr, part: Any,
) -> None:
    """Replace all runs in a paragraph, preserving ``pPr``."""
    _clear_runs(body[idx])
    _inject_runs(body[idx], instr['runs'], part)


def _op_update_cell(
    body: Body, idx: int, instr: Instr, part: Any,
) -> None:
    """Update a specific table cell's content."""
    tc = _get_cell_el(body[idx], instr['row'], instr['col'])
    paras = tc.findall(qn('w:p'))
    for para in paras[1:]:
        tc.remove(para)
    if paras:
        _clear_runs(paras[0])
        _inject_runs(paras[0], instr['runs'], part)
    else:
        new_p = OxmlElement('w:p')
        _inject_runs(new_p, instr['runs'], part)
        tc.append(new_p)


def _op_rename_heading(
    body: Body, idx: int, instr: Instr, part: Any,
) -> None:
    """Change heading text, preserving formatting."""
    el = body[idx]
    runs = el.findall(qn('w:r'))
    for run in runs:
        t_el = run.find(qn('w:t'))
        if t_el is not None:
            t_el.text = instr['text']
            break
    if len(runs) > 1:
        first_t = runs[0].find(qn('w:t'))
        if first_t is not None:
            first_t.text = instr['text']
        for run in runs[1:]:
            el.remove(run)


def _op_add_row(
    body: Body, idx: int, instr: Instr, part: Any,
) -> None:
    """Add a row to a table after a specified row.

    Clones ``trPr`` and ``tcPr`` from the reference row.
    """
    tbl_el = body[idx]
    trs = tbl_el.findall(qn('w:tr'))
    after_row = instr['after_row']
    if after_row >= len(trs):
        raise ValueError(
            f'after_row {after_row} out of range '
            f'(table has {len(trs)} rows)')
    ref_tr = trs[after_row]
    new_tr = OxmlElement('w:tr')
    src_trpr = ref_tr.find(qn('w:trPr'))
    if src_trpr is not None:
        new_tr.append(deepcopy(src_trpr))
    ref_tcs = ref_tr.findall(qn('w:tc'))
    cells_data = instr.get('cells', [])
    for col_idx, ref_tc in enumerate(ref_tcs):
        tc = OxmlElement('w:tc')
        src_tcpr = ref_tc.find(qn('w:tcPr'))
        if src_tcpr is not None:
            tc.append(deepcopy(src_tcpr))
        new_p = OxmlElement('w:p')
        if col_idx < len(cells_data):
            _inject_runs(
                new_p, cells_data[col_idx].get('runs', []), part,
            )
        tc.append(new_p)
        new_tr.append(tc)
    ref_tr.addnext(new_tr)


def _op_delete_row(
    body: Body, idx: int, instr: Instr, part: Any,
) -> None:
    """Delete a row from a table.

    Raises:
        ValueError: If row index is out of range.
    """
    tbl_el = body[idx]
    trs = tbl_el.findall(qn('w:tr'))
    row = instr['row']
    if row >= len(trs):
        raise ValueError(
            f'Row {row} out of range (table has {len(trs)} rows)')
    tbl_el.remove(trs[row])


# ---------------------------------------------------------------------------
# Structural ops (change body length)
# ---------------------------------------------------------------------------


def _op_delete(
    body: Body, idx: int, instr: Instr, part: Any,
) -> None:
    """Remove element from body."""
    body.remove(body[idx])


def _op_add_after(
    body: Body, idx: int, instr: Instr, part: Any,
) -> None:
    """Insert a new paragraph after ``body[idx]``."""
    ref = body[idx]
    clone_from = instr.get('clone_style_from')
    if clone_from is not None:
        new_p = deepcopy(body[clone_from])
        _clear_runs(new_p)
    else:
        new_p = OxmlElement('w:p')
    _inject_runs(new_p, instr['runs'], part)
    ref.addnext(new_p)


def _op_add_table_after(
    body: Body, idx: int, instr: Instr, part: Any,
) -> None:
    """Insert a new table after ``body[idx]``."""
    ref = body[idx]
    clone_from = instr.get('clone_style_from')
    new_tbl = OxmlElement('w:tbl')
    if clone_from is not None:
        src_tpr = body[clone_from].find(qn('w:tblPr'))
        if src_tpr is not None:
            new_tbl.append(deepcopy(src_tpr))
    num_cols = len(instr['rows'][0]) if instr['rows'] else 0
    if num_cols:
        grid = OxmlElement('w:tblGrid')
        for _ in range(num_cols):
            grid.append(OxmlElement('w:gridCol'))
        new_tbl.append(grid)
    for row_data in instr['rows']:
        tr = OxmlElement('w:tr')
        for cell_data in row_data:
            tc = OxmlElement('w:tc')
            new_p = OxmlElement('w:p')
            _inject_runs(new_p, cell_data.get('runs', []), part)
            tc.append(new_p)
            tr.append(tc)
        new_tbl.append(tr)
    ref.addnext(new_tbl)


def _op_move(
    body: Body, idx: int, instr: Instr, part: Any,
) -> None:
    """Move element to after another element."""
    el = body[idx]
    target = body[instr['after']]
    body.remove(el)
    target.addnext(el)


def _op_insert_image(
    body: Body, idx: int, instr: Instr, part: Any,
) -> None:
    """Insert an image in a new paragraph after ``body[idx]``."""
    ref = body[idx]
    image_path = instr['image_path']
    width = Cm(instr['width_cm']) if 'width_cm' in instr else None
    height = Cm(instr['height_cm']) if 'height_cm' in instr else None
    r_id, image = part.get_or_add_image(image_path)
    cx, cy = image.scaled_dimensions(width, height)
    shape_id = _next_shape_id(body)
    filename = os.path.basename(image_path)
    pic = CT_Picture.new(shape_id, filename, r_id, cx, cy)
    inline = CT_Inline.new(cx, cy, shape_id, pic)
    drawing = OxmlElement('w:drawing')
    drawing.append(inline)
    run_el = OxmlElement('w:r')
    run_el.append(drawing)
    new_p = OxmlElement('w:p')
    new_p.append(run_el)
    ref.addnext(new_p)


# ---------------------------------------------------------------------------
# Comment ops (modify comments.xml, not body)
# ---------------------------------------------------------------------------


def _apply_reply_comments(
    doc: Document,
    instructions: List[Instr],
) -> int:
    """Apply ``reply_comment`` instructions to ``comments.xml``.

    Returns:
        Number of replies applied.
    """
    replies = [i for i in instructions if i['op'] == 'reply_comment']
    if not replies:
        return 0
    comments_part = _get_comments_part(doc)
    if comments_part is None:
        raise ValueError(
            'Document has no comments.xml — cannot reply')
    comments_el = comments_part.element
    for instr in replies:
        new_id = _next_comment_id(comments_el)
        comment = OxmlElement('w:comment')
        comment.set(qn('w:id'), str(new_id))
        comment.set(qn('w:author'), instr.get('author', 'AI'))
        comment.set(
            qn('w:date'),
            datetime.utcnow().strftime('%Y-%m-%dT%H:%M:%SZ'),
        )
        para = OxmlElement('w:p')
        run_el = OxmlElement('w:r')
        t_el = OxmlElement('w:t')
        t_el.set(qn('xml:space'), 'preserve')
        t_el.text = instr['text']
        run_el.append(t_el)
        para.append(run_el)
        comment.append(para)
        comments_el.append(comment)
    return len(replies)


# ---------------------------------------------------------------------------
# Execution engine
# ---------------------------------------------------------------------------

_OPS_MODIFY: Dict[str, Any] = {
    'update_text': _op_update_text,
    'update_runs': _op_update_runs,
    'update_cell': _op_update_cell,
    'rename_heading': _op_rename_heading,
    'add_row': _op_add_row,
    'delete_row': _op_delete_row,
}

_OPS_STRUCTURAL: Dict[str, Any] = {
    'delete': _op_delete,
    'add_after': _op_add_after,
    'add_table_after': _op_add_table_after,
    'move': _op_move,
    'insert_image': _op_insert_image,
}


def apply_instructions(
    doc: Document,
    instructions: List[Instr],
) -> int:
    """Apply instructions to *doc* in safe order.

    Args:
        doc: A python-docx ``Document`` instance.
        instructions: List of instruction dicts.

    Returns:
        Number of body instructions applied.
    """
    body = doc.element.body
    part = doc.part

    # Phase 1: modifications (body length unchanged).
    mods = [i for i in instructions if i['op'] in _OPS_MODIFY]
    for instr in mods:
        _OPS_MODIFY[instr['op']](body, instr['idx'], instr, part)

    # Phase 2: structural (sorted by idx descending).
    structs = [i for i in instructions if i['op'] in _OPS_STRUCTURAL]
    structs.sort(key=lambda i: i['idx'], reverse=True)
    for instr in structs:
        _OPS_STRUCTURAL[instr['op']](
            body, instr['idx'], instr, part,
        )

    return len(mods) + len(structs)


def main() -> None:
    """CLI entry point."""
    parser = argparse.ArgumentParser(
        description=(
            'Surgical patch: apply change instructions '
            'to docx in-place'
        ),
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    parser.add_argument('docx', help='Input .docx file')
    parser.add_argument('instructions', help='Instructions JSON file')
    parser.add_argument(
        '-o', '--output',
        help='Output .docx (default: overwrite input)',
    )
    args = parser.parse_args()

    with open(args.instructions, 'r', encoding='utf-8') as f:
        instructions = json.load(f)

    doc = Document(args.docx)
    count = apply_instructions(doc, instructions)
    count += _apply_reply_comments(doc, instructions)
    out_path = args.output or args.docx
    doc.save(out_path)
    print(f'Applied {count} instructions, saved to {out_path}')
    print('Re-extract recommended: python3 scrape.py', out_path)


if __name__ == '__main__':
    main()
